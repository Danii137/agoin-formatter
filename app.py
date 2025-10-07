import streamlit as st
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.shared import OxmlElement, qn
import io
import re
import base64
from io import BytesIO

st.set_page_config(page_title="AGOIN - Formateador Corporativo", page_icon="🏢", layout="wide", initial_sidebar_state="collapsed")

LOGO_BASE64 = "/9j/4AAQSkZJRgABAQAAAQABAAD/4gHYSUNDX1BST0ZJTEUAAQEAAAHIAAAAAAQwAABtbnRyUkdCIFhZWiAH4AABAAEAAAAAAABhY3NwAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAQAA9tYAAQAAAADTLQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAlkZXNjAAAA8AAAACRyWFlaAAABFAAAABRnWFlaAAABKAAAABRiWFlaAAABPAAAABR3dHB0AAABUAAAABRyVFJDAAABZAAAAChnVFJDAAABZAAAAChiVFJDAAABZAAAAChjcHJ0AAABjAAAADxtbHVjAAAAAAAAAAEAAAAMZW5VUwAAAAgAAAAcAHMAUgBHAEJYWVogAAAAAAAAb6IAADj1AAADkFhZWiAAAAAAAABimQAAt4UAABjaWFlaIAAAAAAAACSgAAAPhAAAts9YWVogAAAAAAAA9tYAAQAAAADTLXBhcmEAAAAAAAQAAAACZmYAAPKnAAANWQAAE9AAAApbAAAAAAAAAABtbHVjAAAAAAAAAAEAAAAMZW5VUwAAACAAAAAcAEcAbwBvAGcAbABlACAASQBuAGMALgAgADIAMAAxADb/2wBDAAEBAQEBAQEBAQEBAQEBAQIBAQEBAQIBAQECAgICAgICAgIDAwQDAwMDAwICAwQDAwQEBAQEAgMFBQQEBQQEBAT/2wBDAQEBAQEBAQIBAQIEAwIDBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAT/wAARCAHJAckDASIAAhEBAxEB/8QAHgABAQACAgMBAQAAAAAAAAAAAAoJCwcIAQQGBQP/xABkEAABAwIDBAUCEAgICAoLAAAAAQQFBgcCCBEDCRQhChIVJDFBURYXIiU0NzhEVGFxdneRtbYTMjlkeIG01BgmdISXpMHwGjVDUlhylqEjMzZCVVdiZXWFRUZjZ4KHlaaxxNH/xAAdAQEAAgMBAQEBAAAAAAAAAAAABAUGBwgBAwIJ/8QAQhEBAAIAAwUDCQUGAwkBAAAAAAEEBQYRAgcUITE0QWEDFSRRcYGRsfASFjehwRMXIjI20URy4QgYJidDZHSEk9P/2gAMAwEAAhEDEQA/AJ5wAZg/mOAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAGkgABHPoAAAAAAAAAAAAAAAAAAAAAAAAAAPJmI6gAD2JiegAAAAAAAAAAABla3NGTD+GlnZt/B1DFJI2ntKiXlukjpmixj9qxeJwMN4+/wB5wTNfzJHx87W1ppKzwHCbeMYvWwinPOxMR9epjeuPa25loZWPhLq2zr619QSkO2qaIi7hUi9oqTkIt9rwL1kxdo07rqxffUfDF0PSQcliXuywU/mfouJR1XuWt7ifVIjNr1pGao57phfon8heIyfL+acd5yF4jVrEWp0hkOfsqTkzHowruAATY5RowsAAAAAgAKMt3j0eW62YXYQF0s1jqo7CWgkvXmKoGLaps7t1y21wqiLxSLigGqpjVU41FeL1VRWScnq/KzZityhkeXMs4xmu1weE1Z9vcnlgoGbqqUj4Om4OSqSflHnBQ8RT7PtOSfuvJwLEyf2a3IO8dvM2bycVluqK38S7XnKXjmmVtlYfzB369f1MvpyvZH8r2TuncFO2DszStBriaIzlp5ky7SrOdXDz1fTjnvrvX/tKqaomhzvcO6duLTQDiqrl1/RVuabap1MU7W9Ss6NicHxYnjvEmHzFRx+v8sat94LuGoU63FZlt6eGuiH2nei6535Js3207drLPTi6aq1w1FUsjIp8q9j6fqP2JDotWcrY7LWDvnlgk9uieL6WqaKT7IdqUq1tvtN19b11tmM5m8oGT2zbB1MXoKp2pbkafK8iWbtD5aB38m6rqJ5wTLNRHsNvhw+pxVDa+sqbjsa/G7dQ6YTz9vciFnGQ9zcTws2o1/8AISOXU6P3vJbY7PbvGNnqduzHtU6zt5am4sPIIia+Rk8Rm9dL8SMzFjdKzd2bLTqUxeC1lfWnqbFhRcMXcuj3lNyMgi+CpxRtILJ5rst+Yz8J6SV9rX3O2rNOKdRdF14wqOUZap+M9Y6q8a8/BMSIck3NtLba8NKv6KujQVI3HpeWw9+petqcY1LTb/RPFWTrXCv9nM8jEJjrCLiW5DLmL1eKy1b19+rUnAtHz0dGvszc3BMV9kvqNraas1XE89LSq3b6obaTiomBEwsXnN7FaquLyu2fJETC0TXESNX1y/Xgy03BmLV3vtxUVsK+gkXi4qoWeqSDX4axfJ3J20/PWPIsK9nieUtG5q3f5gyb/FilbXZ9fc4dABOYVHTUAAOkgOR7T2guLfWtqYtna2j5qva3qZ3wkTStPte0ZJ6njqvwRp49+LB8gHRwLU2saw9xM7DxpeKvMSo+ZWrp567j7R0rrhxdZo+dqqPZR1hX8Hz7ozXVUXC7TTEsKxY4XlDNcp5HzFnK1rhfLZ8UlNjcsuYTMdMOYOyFnbj3ZftnnBPFommnklGQP8uf+wmn8+Mr1s+je7xevGuJ9UdN2ls6u3TrKzuZdHtCRXy6a0/2vh1/WX00HRFG26pqJo6g6VpuiqQgWSMoWl6MiGdN01CtkTTubFrhTDh18yctVXwOoF6t5rkGsG5exV1c1tnoWoIzCuF/T1P1P6Nalj9fI9j4rC7eNuflxInL6iv84XNfRW86u5jKOB1eKzLc19/7BKpsuix5vlwKm1v1loXbfB8L6pFwfZOpw7WnRod4TSrXbOKe2thLmrjw9fsijLivGEjg/wBbFKM4nCn1lKe26QrusNi54fY34qVzi8OMaWfrHs79i0OwFst8Nu0ryudgyo/N9a+OkHCdzaXA4202JfkSfZs/qQft7kwkxkPc3bmKlO3Gvr4hr78wWQbN3lf2Trb3uy3XPoOHZr1cVVuoj0S0NhXzLONOLZa/zw6km3XinkFVMPsXrJ/GVHBSrTrNXLVyj+PetlTxRfBfl5+Bg0z4bgrKRmw2MxVNq4RhlqvI6w4nzWrLawuHBQ866RMCYUmqWRUZLr6pVesUaPOa64nWiYVk18Q06sYzFuHsV9njctWuIj2x+jX5g7cZwMlF9ci9xXFBXupfsbbuV4qkasiU7Soe5LVPF5CPV+Xmi99Z+U6jlnE/a6OfbtG7g1zgr0TG16vUAA9R56OT7aWPu/el9LxlnrQXRvC/gmiO5iKtTQkxWslAtsS9VHj5WrN3wuqqiJ8an2Nd5Q8z9radkKxublgv5bGj4tWzN5VlwrQVLRVNMeNXRl390z4PkvJfjKDeivt02N+c1qKiYlS10Fp5cP8AjjEZpukNYcKbri82LqYcOPFUNHKuiaf+tMan9iFXGI+lRUbjwbdhTxPd997ZtfxRry7uTXmH2MdbO4svRM/cuKt3WsjbCmJZtDVHcJnTjyRoeCcvtEZM3057Dau+afWfHFkXRpreUfdvJtm8t/cWDj6jomsLwLTdSQEq07Ri51o9pmOTEip5UVFX+/MkeXsRWnRg2SsuTmnMM4Nsz60boMr29i3YtWbvG8HEwuGTqnLlcR86eWgrZ16pY5PJTM2qonfGCL5OTxnq7RV0es2eKHTw+PwJeztRtRrCnx3ArmXsY80Yv4gAPVbpoAAHUOa7W5bsxl6oaQqKzuXy9l4afjJjsWYlrU2tma2jGDpOasnr1ozd97774fEcKFwfRacOFcneYTFiTXqZmXip8X8WadIVny8Vma7vcrVc45h80XJ0j+yN25uXm+9jmsS8vPYq8VnI6deumMS7uxbeZopZxU8jHi2bTiuS/wC84eLGula4Uw22ydqifjV1U+LXz9zgSOU/VexFqry8DO+XamVcxWsHqTrERH6AAJcdGFAAAAAAAA9jWObym3REXVNVVNE5GwR6Pnk3a5XMlURc2oYxI662afaNrsVJxOiyLOCRMSUqxX4kaPlefLOKnkI8t1pk92+dXN3ae1r9i4dW4i3qXLu86XVEZUzCYtHrFVRF9nriYMddOXaBsv5F/TVvaUlZqWkIynKPo2nnMxLPHWHDGxkNHMGnXxYlREREas2qYlXyeXkUeIWNfRO+XRu4bLUbNi1m+70jpr+b2Kxo2nq9peqKAqqLaTdL1hTzymqiiXeH8JHzMa/Z4mL1n8mLDiVFTkumLxTU1cOeDK3NZRMzN57BTSuXTWhKlxJSEs7xJi9FcE+VX0I8XRfFGj1kip5FRTYAbsPeQULvFre3XrCCiG1NVDbK78pRzylsS+uPYT14r+lJpU/P2eia/DI96nkMRnSbMnfo5tZbjOvR0YjioLSu0tjdLhV0xv6ZfvVWEeKmnvGXfcH8fogxeY8w7Wpa4SerL97GCU86ZM+8eE85r8/d3oxAAXrkKNdNJAADqAGfPo/2722Wae/kjf65cJidWZy6S7V6xbSrHWNq2sUVHzFkiry4Zii8a8wr4rwPLR2unzs2eG5Qu8s4HczHjFXCKfhqyq7kjctxVn4Kks3ea6ndjI3inGbSprW2qqBj1o60zfFzYTE0xVNO19NVZMl5MU5r3vVWVNN1Lw23sXQU7cq7FXQVA2+phlxtQ1VVLxI2LjsKctcS/wD85/Lop/a59yaNs5QdTXKuHMRlOUPRtPOakq2oJPF1Y2OaMk1xYl/Vqprj96PvQLl7wy6Dhu220hR+XejJVzgtVbJHOnJE5TM4iJ3qWfc1/MebT4a9e495CvxVrn0dc4zjOX91eXopU9ONmPf7ZZQ89vSUbk1Y/mKAyOU+luaObdyW9VaQ/adbz2uHCirDQbrucZz63s5HbzTF+KzVNSbS5t3bpXoqLbVddi5laXPq5xhTC6nq+qR5UkhhROSIiOvA+ABcV63CzzcsZjzhmHNNnicUtTEeqJ5AAJmkMTiZ15vZjZJ9DyDeUinzmNkIt5xrN3EvOzZNgZlcnm/izs5Y3cTC1xWDvM7bFti1fUrdWZ7Qqho2TqovA1TpxrVdMKYdX3GM0TwZJ4mGEH4mpErvA8exjAbfF4Vbmv72y6yGb1PK5n9gsLe1su5pO68W1wOatstWmLs2t4XrKiK9ZIqp2ozRVxd9YqqadXiuExL1U5YzsZGLBZ7rWuba3lpNJJ03RxjpK4DXBjiq5oZ5iRMOF5DP+qqp4J1mapwbtETVNDWK0HXdU24qin63oCqpqiq4o2Y7apuq6UmezpOCdfDWL4ppguk33HaZSHtHylvdhIZxGa+hmIr5GbNLZyDVOWGpnrLySqLqqsU7mq6O9dE4Bajy+HzE61HROXN72D5hwi1hOeo56erqwe578ktY5Dr8z9kqxqenKzVsybVNTdUwL7DixzUY+1Rk7fsUVXcU811TgnyroqcuNad9OmB9VXdd1lc2raouBX9VTda1vWUw5m6kqGoXvaMpPul9+vf7+Q+VLer0lzbik1Ld+z5njSnrOgckWctFXt+ri0Tae1tPbeq69uHLrTVPQDdUTC9conNcS690ZsE11fHG5br0dLd6NLLWV2+dC5kFs0ubfeIwtbWs5dn1ZWk6OTF1kepqipxU6qYXi8vYTNlouFXTzWPYsTVnSGRbv8p2s45g817f8saTM+DJLuxN2Lavd0Wq2DNrsY2q73VjENFuldVGCrKTbrrJiwQ0LhVO6xLLqoiJpq8xIrp2uvNPZ3i29Uy6bvqiVWuk2leXfn2uN1QVkqVeolVP+ePCj1+854YpnpovG4kVefdcLtU0PW3rG8bpfd02Ac1dsVb1Fey4j1zTllKHc7RdXz1fZ0w+TXTg4/Doq6664sbJqnVV0a6a6d06+vFW1TXMuXVEzWlbVjLdsVHKT7tO0XznTRE0K+vXm3PF23QWfM90d29KMp5Sj0z2dOn5u++c7e6Zy87D+Xi6xubJW5tg51wtLVWoevKcohWunv5fZkpqqe/tPiRmYyFTRfFF+Q8AutnZ2dmNNno5bxTFsYxi1xWMW5+IAD9aQrYmdddXb/Knn0zTZPJpvLWCvBVtJxHGK9l6JdO/RHbSd1TReNhHXck/lyd9LKt2bv47L5031PWevjEx9jMx8iqNIhpgerittdZxhRMWkG9cr1mzz1OnZ77VVVURo9e8yBdF0XU8bLb7bZbZu4b7fhdu258U056EOxXq2tIZ/lPeVmDJtuP2EzNTXnDarZrsqdoM49opizF7qXb1DR85iV80dpgTDUlKSfPCxmIZ8q6tXbJVxYsOLCn9qYtc3vDsgFzN3xfOTtZXKYqio+c60xaq5zRmkdGXKjE5a+Zq7Y+/mKLqir8Eds3hUpuId7TLZn6a2OUrMTUCS1/6Oiu0LcV9MP8ASTvRBsVwo+avcSr3iWj9FxL1eb1nqqIqtHrzFlO3m2RCjd4Rlrqu1rxY+OuTT+tS2Yqx3s0x+hSdY6rgwri0RUavebJ4ia+pxaqmrbCpTVrHDWtJdA5qy9g+9DKf3jy9pxkc/Hl18h/r+bWVg/Zqqmp2iqjqClKqinMLVFHTDqmKkiJRPXOCdsXvAPmT4/GMljnGrkCdidjb/Z7caTEqi+ise3nms+i6C+1sRm76Q5+S2vP87qO+88cYROise3nms+i6C+1sRm76Q5+S2vP87qO+88cY15btuzLq/KUR+4+17LDXclrXRYvcz5nfp2b/AGIwIpS1rosXuZ8zv07N/sRgWN7nUnVqTcty3h1dPUoSzPZbrX5srN1rYq81O4qioet2eJq8wYMOkhDOuSMpdi78WrtkuuLDiTmionimqYtbnvAsit08gF5Zm0lftcU1T7vWXtZcDYMlwR1yYNFXCr3T3q7TnxrJF0Z+KKqKyersx6zu7b638pbmIrCqo2nX9169w2zoNJNURKhncTKQkeycPPm5VpCP/H/MXTU6qbwLIRajP9Y+oLT3Da9myrbrzluq+YMcEhUtuZ5MOHDgeslVdVbr1U4xlqmF5hVUXmiKV1WxNblLofedkOpnOrNylpxtf8+XTRrAwc8Zmctl0sp13q2szdunFpyraOd+XCqR06094zTF6vspo+8y+BwOZLE6xq4pt1LlG5NG7HP2AAD4BcH0Wn3HWYb9Jd592adIfC4PotPuOsw36S7z7s06VmIdm+LcO47+tKvsn5OFulXe1lk4+fFYfscCRyljXSrvayycfPisP2OBI5STh/Zfh8oRt9H4gWvZHygABKanAAAAAAA7jbvvKbLZ1s0lqbFMMLrYQFT1H21Xco0VEWDphh36ceIvhzadx/lsgyEzokYVUuYveilSjnrCvTo5OSh3YHKZt8wlXxXCV/modNp2IV209cYSjWCKkEir4pxyq8fa/n7M/v0jHOi6sFlPY5dKQlcTS42ad25gpdGjzWQhaOYJhWdxp5uO1ZsNPLx7wz9RDCnLe0tDwkSwjadpCi6cbQ0UzbYsMdGQkZHtFw4ETVURGrNqiaeRNNOSmtQ3qOcLb52s3dz7pMXzhzbiMerbS0LTCiojKmYRVxM3yaqq+uCq/faa8u0Ci8hHE2+fc6wz3iVXIG72tl2j/PZjT/8AdzTuPc6C5QM6NEMqilFb2xvrwtmq7wunvVjGSvnusFM+buLzlz8Gcg9U2FF4LS0dfC09xrP16yWUoq51IvKPqJpixafhGkgzVji9UvvlNU0X4sPyGpiNkZuaM5651sktBVRUsrikrrW62mC013cbpOvJv5OFapwcuvWxKq8e04J6q6YUR3jeIn4up7iHkNJi2x7chmXZvU7eT7vf08de5r5MytiayyyXtuvYmuNgvontNV7qmXjrglwLPNefAPmSfBHzNWL3n/0gcHFafSd8myo8tRnXo5nh4d0jazF6eCa88TjXjaVmPHVVTvrFf9ZkhJYWFby/FVWlc8ZdnK2YrOE909AAExiUcugbO3dhZUdjk8yZ2Ssrjj0jqm9Dfo1ufiw7NUxu6nm9H0gmLFzTuWvBJ4epYtPjNezu77TbC9OdXKjbR837Th6mvtA4qkZqunHxbB52g/TX+RsXxtOdmiN22PCiaJg2ei+VNdCkv850jq6U/wBn7BdjYi3mPa9iNnpNedvbNNtb/I1Q0rtsbDC0b3KvsrJ2mJHnqtaWhHvlTmiTSp/7NjyJJTtTnrvi9zK5ucy96Nu+WTbV1dOUeU3iVfVJAsHvZ8CzVf8AwhiwOqxPr+Q4arrDUG8DMMZnzXaxTZ6Ryj3aAAJkdGFAAAAAHUAA1gAADTl9l2fyV5cXubHNJYmxWwRzw10K8awtSOmSokmygmPf515h10TuMQxfaJ8ZtMICmoKlIKJp2Dj46EgqZiGsLExLLTZx0O1Y4cOFk1RNfIiIn6vLy1hl6MFazYVXnDuZct+2V21s9ZZykU8VdEjpSdeMWKL8fdEm00+Mqz3q95nGXzd8ZrLlRj3DHSra1zmjYV5oncpSq3rOlI95p5Fwu5rCuvxFHiEelaOst0NWrgOR7eZbMfxTEz/8IQa72LOa8zuZw7pXFiZRw5tjSDz0tLSNsGHSOawTB2iccvxv3eF68+JHyJ5DHCAXWzGmzEQ5fxXFLWMXreM2+uvL8gAH6VnWOYAA9AAB9/aa6VUWVuJQN0qClNvDVdbuo2tZU7KNvVIycsfFFTzG0oyk5gaTzS5cbLX7pTRvE3NoRtU6Mk54oF2q+uEQvJObF20es8XxsVNUwXA9GCvO+rDJ7d+zj95xm3sld7j4fXCmLDHRlUtMMhhZ6L5Uesptdfz1CrxGs31uGxy1Vx+cF2p5WOfwYaukZZX2di86UddunorFHUfmbp5axeYmjT1r2tTsO4TvPVeeiwbxVXTVZ1eRgTLgulF2xb1Hk7shdRu24iZtlfZrDq8RPUx8XOQ0hx31vIWE+oh9JWHxpUiWGb2cGnC88WYqctefx0VF9FY9vPNZ9F0F9rYjN30hz8ltef53Ud9544widFY9vPNZ9F0F9rYjN30hz8ltef53Ud9544pPLds2W68pfgda9lhruS1rosPLLPmdX/37N/sRgRSlrXRYueWfM79Ozf7EYFhe7HOv1yaj3LfiHU19TkHpPsg+h8lVgZeJfOo6Rjs31PvWUszedmyMc5SmKw0eJ8nNf1nIW5I3tUZnRoPYWAvtObLYZp7exWFMEm9TDH7O+kWy09eGP/ejHmr5kvPnxWHXV3hZca9KQ0w5HLJYcK8sOaaKX/7NrIibt/casLUVrSFxbd1TI0XXtDTDapadqCLdayME6ZeKKh8K9fiq2vc2JnPO13Jm83iPI86k6ax7ubYn72XdhUTvErL4sTHZ7CkMw1um7t/aC4W3TCvG+p63oanETmsTIKip4LiZYsXF804xm812Vf22rG19b1dbm4lLyNF1/Qsw5pmo6flmvrlBumPgqKbEbdPb0SiN4tZ1q1llYUpmPt00asLu29b7TVX3LlU0Hy1WJfKnPX2E8XE1VV7m8edZ99vujmOdGhHGYCw0I3b5r7exKrji2SpHYL7QbJFVYZ8mv+NGSKnBPU8/Crrq0xMlezNT0O2t94eS6OfsIjN2U+1cpnTv9aCAHsyUa+jX8hFSrFzGyEW74J40ds+zZNg6PWLuJjTk5P2tna2Nr7O316e8Lg+i0+46zDfpLvPuzTpD4XB9Fp9x1mG/SXefdmnStxDsvxbg3Hf1pV9k/Jwt0q72ssnHz4rD9jgSOUsa6Vd7WWTj58Vh+xwJHKSaHZfh8oRt9H4gWv8ALHygABKanAAAAAed3ILZ+jQ5N0ttZKus3dYROBawvm7WjrcYneFFkYukYR4iPXfjyV/LM115fiwbHmSP5Q8uVU5t8wtobB0vxTaQudV7WFeSrT1WKBi079OTKfGxZsXrzTyrobTG2tB0vaygqJtjREW3haRtzSTOjaeiGqJ63RbBphZMcPh5MLJE/UVeIWO5v3cLlObeJzmO3HKv0YVukFZyWuVzJVIWzp572bdfNPtHNp6cRu4RJFnB6ItVP0XyojR9hZfFinE8xr8NFTxTTyl1G9J3LOYLeI5h1vHhzL0XRNE01STSjLdW+lKNdyKwTVNHj/Hixaoiu3rp49VV110Rj5WiIuPVeir3rXn/AAubc4vjS3Ev++irZq1q3PqtN6GUs85szBxFOr6HHRLCqqq6qZzuj65yly05yY611UyfCW6zO8LbF2rpz1o9jU2qLSrtPldvHzH5Z9F8h3OToqN5EXVM3Nu+Xl9LeX0T+vH9GnRY74NXTd+wziW/byEW741k9Z0FLpJMXPnTvh+bNmrZ0hiuWsh7wcAxiri9bDOz9fGPirOzb2ApDNNlwvBl+rbDrT90KLc02rlGibTFCO10VhL4dVRFVi74J4mq6ask8eZqybo27qi0VfVtayvIzDDVrbKr3VHVFGYcSLgZuWL3gXqovm8qL8ZtdrYsa3h7e0RGXKnISpbjxdIxkfX0/AMuy4qflOE0fvmTJfYqPXa41TD4c+WmiaRjdJkyaY6BvBQOcmkIpG9L3hbel/dHE1TXg6nYMtWD1eXi+iGK4dPH+L+vlI+HeXiJ0bM34ZanFcIq5krR/FX6+9MUAC+cnRznmy3bhbZ7F1vRMqGxe6fgFeVi8w6+GrK39XJz+tDYqXCeOI2hK1lGCdaQjaSlHjVdfByjLEuFfr0NbFuaq8W3u8ayZTm12ybDYyl0Vt9hRV01WdYyECifr7cNmc9ZbB612zJ5sUcN3TRWrr/Nxapz1+so8RnS1H16nWm4nScm3an/AHE/JqDAckXltxKWmupde1k2rhZe2Fx5O38urtdU4qCeyDDX+onG5d7PSHK1/Zmr5fbiY75+YAD1D6gAE8o1frY2ftbcbPrcwxuW++0wwj5WKsReObh5RnxsPLRNt5mSi5Br8NY9zPf/AILmYT/R5vZ/RVNfuZsxt36rNcimTFNkmz4hMplueWPTx9BsGi/2/wC87gYcCc126skTzL1UKXjdrXo6Ww7cFh9qns25xXrHg1PP8FzMJ/o83s/oqmv3MfwXMwn+jzez+iqa/czbEfg9n5F2GiebFh0CbLZ4tdF2K+fTEijjvWk/7vGH92Kz+TU7/wAFzMJ/o83s/oqmv3M4XkY19DyEhFSrFzGy8U84J5EyrPs2SYOvgT4282LAm1VergwcGuDEvLB4aIapLOLUmwq/NNmoqvY7fu9U5hKyqZmv8vqd+/JtexxMxDW28XdxT3f1atripscROny/upZ6J+2b7SRz3PF5P2mG3WwbYsXgiO1rz9yT6lMonSK5HbtN2RXDdvz2ErcmjWL5PO2Sa45f2AxCdFSrDYRV184tvF22HjqxoSmK1bt8XJcGGBeSLFV+uqMJnH38VBvbg7sDMTsWGwRzJ0KlP3AbNkTRVbMKmj+O1/mavSBY5YrEtx5T0tbm7OzVjnw1hrlQAXjkSNdNJAAAAAAAAjnOgVq9E/ebbDN552K6cE5aW6d4NfKvVrBdP66SVFkfRUaDfRtrc4d1NvskSOrG5FMW+ZuVTnjdQTGQfvk+qqmRExHskz7Wz9ztf/mFVmPV+jIV0iZi3dbsC6G32v8Ax8XXlHvmiaeXFNx7H/8AeU16X+V/v5y+bpLdabClN3Ywpv8AD6be41+qYplm18/AtJGfX7DIGf8AK/385Fwzsa938TE56/8AWVF9FY9vPNZ9F0F9rYjN30hz8ltef53Ud9544widFY9vPNZ9F0F9rYjN30hz8ltef53Ud9544r/Lds2Wxco/gba9lhruS1nosfuZsz306t/sRgRTFrPRY/cy5nvp1b/YjAsb3Y5al3K/iJU9j7vpRfuG7I/pSRX3NrEhrLlOlF+4bsj+lJFfc2sSGsYbOlWPr1JG/CZ++tqPCPlDnjLPmSunlOu9RV7rO1FipytaOd8ZrhxKkdOtF9nQr5mnN00ffEvJUNkDkBz2Wo3hFiqdvFbyaxRc+x2uGJuNQbjbphlqEnNMOHFH7RddMTTEqYlZvkwojzD5MOJHbU1gx3VyEZ5bp5Dr1w92Lcucb+HdKkPcegNu/WKj7kwaL1kZvV+F68mT1dUZ80VFRVQ+1mtxXRC3X7x7WULnA3+xz18FJG/k3QTau42ezs5aafxLXUc3xzOYC2kBsOutYNueL0YsWa6LxrLCmPE+wYU0dphV0mjpHfGxxfm/k8xtW8r2ZW1ubu0FJXvs1UKVHQtStE6mDGqYZKDd66vYiYZpzau2SqmHFhxac08yp1pAt+7uiEsbN1BnAy3UsuGy9VS3HXeoCBZonpNSj5V1mWCJz7JfO+a8kRk9xeVm77lBoWJj0a2zzerkKpfpxnnLfZNOcR4/9dNCXB9Fp9x1mG/SXefdmnSHwuD6LT7jrMN+ku8+7NOknEOy/FiW4/8ArSrHhPycLdKu9rLJx8+Kw/Y4EjlLGulXe1lk4+fFYfscCRykmh2X4fKEbfR+IFr/ACx8oAASmpwAAjrzADmDL5ZWsMx147YWQoHYcTV92ava0ZEL4pHo+9nvX35owZ6vv5gJnTnKRTqzeuxRpd+ir3ox+TP0PURc/PBWEX1ZCvVc2ltAjtlqixTB7pPTKfy92xYskVF5dgPPHi9FzyZ5N4FYjd/UTR9a37f1ErKsKuWj6ap+h4n0R1K+dI1V86dqyxOmvdGeFFw48XNU41lonNDsRl8s5R+XyzNubJUIx2jWl7UUi1o6L2ePHhxYniMcODCrt5oq6OXiqrzF/wBp9iXykCW/eznLmnzm1RStLSuKQtTltxO7UUjhbPVWMfSqvP41zGmqoi8WwRivV5K0gWa+VTG9OJtaz0h2HjGJRut3fVatTTjJ/OZUj/4S3u6Pgl/f6LGn76P8Jb3dHwS/v9FjT99IHAWPm+r9Q03+/LO3qhfH/hLe7o+CX9/osafvoTpLW7oTwZ38T/5WNP30gcA83VekfJ7G/LO/TSGxOyv79HI5m4vHRViqDfXIhq2rp07aUx6YFHtKdp6QdsWWN+rJXuF461dIjJNMKc9fj0O52f3KrB50crF28uco4aNZesacxvKFlXWnWg55gqPoJ75V0R4yRF6vNWvFprpqawyg61qO29Y0hX9Hyjmm6voKpGtaU3ULVPXKClGD3j2L3+pIbRTI9mipzOJlitHmEgF2DRrcKksLmfiW+0TH6FZ1kvATkP50Vi9Zv8HPx6mvlK6xX4SdY7m2d3mfP3gU7WXcwxH2pj46tW3UdNzlHz1Q0pUkU6hKopiYdUzUkTLf4zgXTB7wD5kfjFBvSNMmuwsNmlj8wtDxf4CgMzjZ28mcLVU7MgaxYL6+aonw9FYPv5Yr3zk+RkFWYtexzFmfBbOXsYtYRZ9fL9H1NCVrO29rGjrgU5t1bVPQlSNa0pt38AdsHqP2P7CbW2w936bvxZq1V5qO234amroULF1nC40RMWPE2kGKPUbfqRdFX4lNTSWVdGhzysKkoKp8ilfSeFKooVXlwrLYnTnD69wb52r2ciGfNFVWDt4jvlqqpIu/BGildiNbSdWz9xuY4wvHtrB7nSwxl9ItymPrG5x/TuhWKtrc5nYf0Sq7bYkWPY1PH8BHzjP/AOJEYvv/ADB6YCdNP/ybP7eO5JqHz85WazsnUblpC1BgVKmtdWjpou0Wg6nYKvAPkTCirwyorxm9w6Kqsn7xPLqa0q8loa+sTcStrSXLp2RpOt6El+x6jgJXmjHx5ovvpnzTvqDDrHdKv3xZRuYBmCcYpx6HY/0cbAAtWnPAAAOjkBneC50a1bsWFza/jmEWz4FkzZ1e97MYNT2PTtu3/wBblyP9vHv74cbg80j1JPHYjp9mNqfjLkj07bt/9blyP9vHv74ZntwPH3RvjvCbavZav68mqPs/Tk9derWMvUryQjZLuKwTFERfFe1ptli080epgZL7ej9ZBpnKdlkfXiuXCbWDvLmTxs6mfQ7xmqytLUwx63YbF7y5PHvGvHz1dVRVftMKonB6kTEIitVhszdNhWIZgzVWna2p+zW9rM3mKuzF2GsdeW80vtdhiY2ntXPXAXicSYcL1I9g/fq1Ty+LJMPLzoan948271y4evtu4cuHTzjXjt34qXPdJNzc7OzeUyAyzwsphSt8ztQIxfM2TrCsnH0xBvML9871111ePEhWXgiKm0eeOikLJGw+NIZLv7xqLuP1MH2f8PDLpuJcweysFvC7MLJvezqavY0c2AqN1qiLiWc5wWHVfDWXYwmH9ZsM702upy81pbo2hqhExU1dig5W39RIiKqq0nGb1htERPP31dPj0NTFGSD6IfMJaKfuYyYinnHM5Zo89cmLr4abNzdhZ2IDPplPoC7+weNMFwotilHXop5rjT1iqdh1cMh1cPijR76h8y5acHIYOevh8L/kJ1i3DJdxGY6t2layfd6z0j2tbDeK1lT2PuVce0FXs0jqut3V8pRtRN8OLXDicsHnA6ovwReWnynG5YB0i3dqy1abDHn5stTjqSfQUQ3h8xtPRTVMUg8i4/kxrJPKvANdWT5OejLg3fJGbvWP8n1rE2dIlo7PGU7WVsw2cLs/yzrMT4SAAmsR005AADzWAAB7rz8Q2We56ysPso+Qix9v6kisUXcCpYn0zbkNHaaSjKUnerIKye/nTFmrFiv/AIeSMbizdxSOcG/cdea4UI6TLzYmpGkvLPHLREjbjVOxXj2VNIioqO2iasXz3kqcGqNF0R4ql6d27t0NYq19b3fuXKNqeoW3lNuqlqOYdYets2DRlhVVVUVebjwTl5V/WUeIWJmeDh09uRypFClazli/u18EfPSlL+bCZudlzy0xEgvUt3TkndirWLTF6jipxcTGDxdXRNMTJrCyGvnSdwksx2HzbZjKqzY5hbr38qnim0hc+r3U0zild84KL9gwcOvxsWjJizVfzA68FhX8hw1Xk0dnjMMZozJaxaOkco9kKi+ise3nms+i6C+1sRm76Q5+S2vP87qO+88cYROise3nms+i6C+1sRm76Q5+S2vP87qO+88cU3lu2bLoXKX4HWvZYa7ktZ6LH7mXM99Orf7EYEUxaz0WP3MuZ76dW/2IwLG92OWpNyv4iVPY+76UX7huyP6UkV9zaxIay5TpRfuG7I/pSRX3NrEhrPMN7LH16kjfh/W1r2R8oAAWrUGmvJlO3UO81rDd6XkTbyiyVW5eLivWjS71vmeqLhVF/wCU8H4ayzDze/k7n8CestiPSlV2szC2ih6op19TtxLVXMpPC8ZO8OBZWmqri3zJPeaouqYk1RWnmVfBddNTUZ3NyrvaXGS+sthYe+E05ksqVxJnXC5e4lfJYiVe6qkyzXRdIh8vs5lounsxOaPOMqsRr97ee6reL5msfdrGuxTy59Of6PkN8huopTIhcDDdS1sVJSWVq58si086TDx+K0kouulNPnaqqozXRVYvMXPTkvNn3zOb0XDVcm2YfBz9VmOdry8eVM06pQPdK2dqszdnZ23dw4WFru2Vzqc4F6zR5s5RhLtHmmJo7ZvE1RV16rpo8wqvkVNF5nQXdO5B5rd3UrmVs68qNrVlEzuYR3XtrqrXFhWWmoN/BU+zTZyOmujxi7YvmfL/ADU8UVCv4iLFbn1bZw3d3GXN4FbMmD9inX84YkulXe1lk4+fFYfscCRyljfSr/azyc/PmsP2OBI5C7odl+HyhovfRz3gWp8I+UAAJTU4AAcu8KyejFZNmz2cudncrOLRW1LcTZyy/GtE6qOlXSqJllz1RVbcCyRURU74+JbbZ25qi71xaBtbRMZhma1udV7Wjadi8XqcL1y+ecCyVV+Lzm0nyk5d6PyqZbLRZeqQXBjirXUg3h3klga9Tt2SVcT6amF18FfvHj54vk6z5fkKvELHONG8tx2VNrFMfnGbnZK3T2ure9yzot8j2SO49xYeW7NubWLZLZWfTxkO3pvVEd4fNwDXC9ff+X/VrRdttttttp+H2/shfAz2dIZzk/wi83a2apeU4i2OVniqKThHaLHzlXvU0ql3+Ki9w4JlDaLrorJ/z0dmBRVVeakjDq3o0wpt8Wa5x/NfA059DrcngAExqbWQAB7rIVF9GUzoLSlzri5L6xk8OKn7s4nVzLRI7d+qYTrBlrNs/Dmr2KYo95qiJ2DiX32S6HIFqLm1fZa4tAXZoqUSGq62NWtayp1ziTXCrli845UVPrINmvNrWYZZk/MUZXzHUxae+YiWyO3seTXZ54MmFz7SR8fhc17CN0uba11rriSpoNMWJi1RE1Xv7bG9YL5ESQVTWc7fY7dntXDfbt3DZw17k8aO0NrJlczB0hmly+WhzBUTt8O0p67FGtamRpidpI4oJ0uvaEPjVEROKYu+NZron4zBfMhB/v8AHJriyvZzKhramIlWlq8zmJ3dmnFatPW1jOpi/jTEa6J4u3qP+Scln2aeQg4fYmPRG8992W61+lVzjS9+nj9fJhTOSbQ3Zrex9waQunbqccU5W1CVG1qWnHWJOTJyvLmnwQ42Bd6RPKXNGxtbVLbi9szz7my73bu8WtlvH7Gta2p3bN6buvS2JrD3etUs1pI0pJoirxTNEXrYol8iLiZPfLzTXi2uJDjDecbpyzu8gopZ6PfY7Y5haOZ42dCXYaMFwY37bCuiwtT4NO+McSpiVMWvGM1VFw64cTxm8gLyzZnbwZUbn0/eeyVT7ejK2gMWFt1MSdeKqtthxJieQ82x99NHqoiovkVEUu/3a++sy655GETQlXPo+yuZRO6OrYVXM4ey65cp75pR8unF6+Rjrxia6d8REeLSeXrzWnWo6wyZnzLufsH+7mbtOL6fUoZ80WTi/WTKt9vRF9baSFFPkxaU7UGLCsjQ9dNtE75DTfsRyvPmi809+czrAbam51nrY3mo2QoS7NA0pdGkJTDo8paq4dpUkXj+XC65L8viTwZlOjKZTbjyMhN2EubWmW+WdqjtIBy0W7dtmCqnNGbB07aPE8PHjdPkPvXxDTlLCsybj8Uq2eMy16THVDqCj6u+i7504d45WgLxZeq9hcKas8cxNTND1K+/mPZDtn9bw44ZdGh3i7xz+C27mwMa3X309uM9xYvHzNWWp9+Kqa9/1o15+7jO8Tzwv6+LAQeyzZbZ47bMWOxcunDp3wTNo0TTjypO0/RXruvnbdxfzNLQVJsGvN3E2dpB7WsjIp50kHfZKNPkVm8KDMlu6ByQZHnjepLc24xVrdZmvV9OG6z1a1rhn4LqxVURlFaaf+jmjRF1XxPjOI1WQ4JuTzZitrW96PssEO6B3FE8tQUbmlzu0bih2UG8bVLazL7UDPSRmHGFUVnN1Y00VGrZExJ1YJeeJdOM9SmJm7q2vre63WW22FW3dulUcbTdC0TDrOT8i6RFTEiLi7ozwKurp08XHgaNGiaKqrhTVVVMK/KZoc2tgsnVuH1zL+V5EUXTLdOpEtHmPCtR1Q6TDyZwzBMXGO3a6ImmBF/GTwTVUgI3pO9gufvE6y2FORTSStxlwo6XR7Qls8T3qPpx0iKvbNUJ76d/jKic2bLXTVdHbx5X6WbPPa6NxYhjOVd02X5wnCZ+1en3zMuuGf7OfWOe7MfX976i2DyOhpTEtM24pV6+XHioSmGSqrFlz+VXr1PBHr96icjpgAZHEaRo5DxTFLuM3pvXOuryiqi6oZNN1TvGai3fd82lSSiyc1ZG4atadvVRLJdUfNk14GZZp4K6YccvJfM8a8ldoqYygJjWNErCcVt4LcjGKfa/r9G2st5Xttb627p+vaBqGErq3Fd0921TtQxGJJCnqgjXqIqYsC+GJFTl5PLqnihHNvZNwbVlASdTZicjlKvKutxJPnMtXuX+n9msjUdDKuHrcZSzJERHbPwTg+Ttkq6JxjTkyxtbrre53U3eVTYaZlNlJ3Oy01JMcXV9s3L5FkoJziXDiWYpZF1Rq7XqoisvYbzTxTXjWd8GVjOBl6zo259MvL7X8LWsNi/4OYim2JI2paVdKidzmmHsxo7TnyVPIvj5aTSzV57PR1jh2JZS3tYPweL6Re+E6tVlt9htme1cbDb7Bw3cNk4J40d8tDwbLHObugckGdp3IVFcS2Tqi7nyvqXd4LTvVoquHnivflRFZSmv/eDR4nJNCfq6/RXrkM3O2c2JzS0rULF2nc4i8VHvKMlI7T/Ofs+MR2vxIzaEiMSqzHi09je5PNeFTrRjiY7kqYKGNr0Y/P7gcpsNhcDLM52GNNUdN64qbBH6fIsNqc2W46KxmHlHWx9ODM9Z6i2KYdXXpaU1M3Ff418nqnaRP9/ISfONTvljlfdfna1taRhmnwS9+Bme3aW5cvjnjlKeuJcBjN2ey1K6R46uBKsMUfUtym2LkjOlmLlF4lNde/qnBJzXvunBFRmUzcFZDssjmPqmpKWlMwVfRio7az14Eaz9OsHOqoqsaZbIjPRU05PUdqipyxGYa41ybdWQoqZr+5dX03QNv6ZY8dM1VVcs0p2nINsmnNca8k5eCL+pSusYhr2NtnKO5HYw+PO2cZ6d393x1j7F2my52qpG0lpqcj6MtvRsUkfTsZE9XDhxaJs1V27XVeJc4sWHisTvFzVVXzarFhv3d661zOVK6ym5faiR/YC3cxxlxaziX2sXeOcYrqjFpi10eRTBVXw9S7eLxWH2GzxYvb3se/hqDMjGVVl2yj7Z9RViZPEsLVt19mixtcXmaJr3Njrziolearp315qidz0etFnEJGH15meLtqzedvNqeVqfdHKUeidJmOXugABadXPez15qi+ise3nms+i6C+1sRm76Q5+S2vP87qO+88cYROise3nms+i6C+1sRm76Q5+S2vP87qO+88cY15btmy6xyn+B9qPCw13Jaz0WP3MuZ76dW/2IwIpi1nosfuZcz306t/sRgWN6JipOrUm5X8RKnsfd9KL9w3ZH9KSK+5tYkNZcp0ov3Ddkf0pIr7m1iQ1nmHdliPruSN+H9bWo8I+UAALVqAAAexPOJlUXuGd73s7eu6YyV5kaoxYrbybtISxVyKheJiw0G56yaU09eqvKKerzYr7y8PYnV4K3LDiw4sS6J/zUXXyKimoCLWdxNvf9pe2Dp3JZmQn8S3ppxqjKy9xZ54mt2Ytjh1SGfvF01lmKctffjPq6ortFV3RYhW/xUOoN0G8f7f8AwnjM6admmfk+B6Vh7WeTn581h+xwJHIWN9Kw9rTJynlSuaw1TzdzgiOQs8P14Xn4fKGtd9Ex9/7U+EfKAAEpqgAAHeDd65w6dyOZgYm/k5ZlvfOYo+mpRjR0A5rxKIj6SlH+jBZpH3BvOK7mr5l/PzOlUnSoqwlKbno+jMm0PSFTOId0yp2oHd6sNSxcK5VmvBvMTJYVpxaMlVFVkmiqqJzJTQQuGqx1nmzPB8+ZpwGlNHCLOkT3fUP0ZiYfT0pITks+dSUvKPHU1MS7t52lJv3T7xevj84Amxy5Qw7a2trb2p29qecgAD8gAAAADNdu0t9nXO7ytFWVkpCy+xvnQ8tVi1hSLNxcn0DpQzl8z78z14J2rpo9VkzeoqKnPE98eLTT9jeRb52nd4nZmJtNUWUXYW3n6QrBrW1JXAbXj9E8hALoke+Z8B2O04lo/Zu/xcLxE1Vouiq11MHIIcVasTznmy3az9mnawH7uzY1p/H8wAEzoxLrHMAA0ieUvdidrZ2o2tidJZfMp2/Qz2ZV28fT+3r5vfu3Ub3NrSl9cb2o5Jk1XqJq0m0XtnrabPCiKrx40wp4Mk1M6doelK5cal2WLYX3y63YthLYlRHLu30vEXXpzBy0xKuJ0sQ9VdeaJhZuvlIsARLGH1e+WxcG3pZ2wGIp1LOuz482xKpbf97rCpUVxIZhpak9viwdVWFUWirHD8WqYmjN001+PU5Ceb9TdUtNj+GcZtKew6fBbc1lIa/VEGtzBF83VWWRv7zdHPha3xlfncjpJO7conZbb0LTF2LyOPe7S3ltsUYuL5PRBiiNP95iKzEdKIvlVbZ9B5YLJUlZ1s46zFK2uW79MatV5+peMmHdGTVUTl37jE5qunmmBBKjDqkdIUmIb6c64prHE8NHh3uWb139vDmKq2Qr+9NzKuujW8pz7WquX7R7PTl3Jj8EafmLE4mAJcREdGsLdu5dnjb86yAAI/UAAPaHJ9oL13asHVjC4ll7mVXbCt41Oq0qCipfs6SX8xf/AAtn+YnGACRTt7VHai7RnSVO2WzpPt/6Oax9OZoLM0lepg1RGXo3t489LiuE87x8w72ydLovvHg9PHTz5hrb9JK3bldtdktYS92bOuFTV00uHap7KKn+z6y39hAGCJ5upz1htDDt8OdsKjtPEx6vg2RrXfp7qnb7H8Nsc2tOL5+KtxWMcn++HOGqz6Q9uvqOjsLiJuzXFz5Fth6uOIoC0dTfhdPjxyrSKZrp/rGvXPOq/F9SEXzZTXdjf5m6OlWsrazBdKcVMEnE5WMs22bbZVTga1zBSy4cbDz9al4pef8A9YQnMzSZ2czOcqd2FRX9uzUdfI1ecdEUmuLsyh6UVERO5QjXubZfDv3ip1XBKr16tblDAcw7wc2Zoj7OLWuXqjlAACWwrx7wAAZNt1zvL9vu1a4udWTez+wvSlzaSa00sThrtKK7D4F4r7jF1ZPOK8fJovLxO7G8B3/D7PJlorjLg4ysbC12wrOVipVKq9OBK3Rh2FNsJDXguxma69x08SfEELhasTrEsrr52zFUwn7uVLPof9wzTbrbfKvt3BbW41u9jl62F429xq8W4CSnpq+gjDB6skY8Eqdju9dOB5fKYWTzhVcK6pp+tNSTZrcTporsExrFsv3PPGEdrZtd57vrH28ZsnSFl9vltb2cb0bdNrdlKgW6iVsj/gYWfgkZcF2Oz/6c1/mJhJPOJVxLqun6k0PArVuGjm8xzG8VzDb874t2sAB9FNz7wAAD9GHmZWm5SHnIOVkYSoYF41moeWiXnZ0nAumPsB6xfH5wD9bO3tbG1G3szzjoyeZ696TXufm0eXK3d06JjWlcWLWTe1JcyLmfUXVcvWMfHo8SC4PubruSryeaqqr7D5aYwwBEacoTsRxO5i1uLmLTrMd/ujQAAV4AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAE8+oAAAAAAAAAAAABHLoAAAAAAAAAAAAAAAAAAGs9QAAAAAAAAAA69QAAf/9k="

# CSS MEJORADO - Degradado desde inferior derecha (verde) a superior izquierda (blanco)
st.markdown("""
<style>
    .stApp {
        background: linear-gradient(to top left, #1a5c4d 0%, #2d8b73 20%, #ffffff 100%);
    }
    .main-header {
        background: rgba(255, 255, 255, 0.95);
        backdrop-filter: blur(10px);
        padding: 2.5rem;
        border-radius: 15px;
        border: 1px solid rgba(26, 92, 77, 0.2);
        box-shadow: 0 8px 32px rgba(26, 92, 77, 0.15);
        text-align: center;
        margin-bottom: 2rem;
    }
    .main-header h1 { color: #1a5c4d; font-weight: 700; margin-bottom: 0.5rem; }
    .main-header p { color: #2d8b73; font-size: 1.1rem; }
    .info-box {
        background: rgba(255, 255, 255, 0.9);
        border-left: 4px solid #1a5c4d;
        padding: 1.5rem;
        border-radius: 10px;
        margin: 1rem 0;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
    }
    .info-box h4 { color: #1a5c4d; margin-bottom: 0.8rem; font-weight: 600; }
    .stButton>button {
        background: linear-gradient(135deg, #1a5c4d 0%, #2d8b73 100%);
        color: white;
        font-weight: 600;
        border-radius: 10px;
        padding: 0.75rem 2rem;
        box-shadow: 0 4px 15px rgba(26, 92, 77, 0.3);
        transition: all 0.3s ease;
    }
    .stButton>button:hover {
        background: linear-gradient(135deg, #2d8b73 0%, #1a5c4d 100%);
        box-shadow: 0 6px 20px rgba(26, 92, 77, 0.4);
        transform: translateY(-2px);
    }
    .feature-badge {
        background: linear-gradient(135deg, #1a5c4d 0%, #2d8b73 100%);
        color: white;
        padding: 0.4rem 0.8rem;
        border-radius: 20px;
        font-size: 0.85rem;
        display: inline-block;
        margin: 0.3rem;
    }
    .stat-box {
        background: rgba(255, 255, 255, 0.95);
        border-radius: 10px;
        padding: 1.5rem;
        text-align: center;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
        margin: 1rem 0;
    }
    .stat-number { font-size: 2rem; font-weight: 700; color: #1a5c4d; }
</style>
""", unsafe_allow_html=True)

logo_html = f'<img src="data:image/jpeg;base64,{LOGO_BASE64}" width="80" style="filter: drop-shadow(0 4px 8px rgba(26, 92, 77, 0.3));">'
st.markdown(f"""
<div class="main-header">
{logo_html}
<h1>📄 AGOIN - Formateador Corporativo</h1>
<p>Sistema profesional con formato oficial AGOIN</p>
<div>
<span class="feature-badge">✓ Títulos verdes</span>
<span class="feature-badge">✓ Alineado derecha</span>
<span class="feature-badge">✓ Logo en pie</span>
</div>
</div>
""", unsafe_allow_html=True)

def add_green_header_paragraph(paragraph, text):
    """Añade fondo verde y texto blanco - ALINEADO DERECHA"""
    run = paragraph.add_run(text)
    run.font.name = 'Century Gothic'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), '1a5c4d')
    paragraph._element.get_or_add_pPr().append(shading)
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)

def extract_project_info(doc):
    """Extrae información del proyecto"""
    info = {'title': '', 'location': ''}
    if hasattr(doc, 'paragraphs'):
        for para in doc.paragraphs[:5]:
            texto = para.text.strip()
            if texto and (texto.isupper() or any(kw in texto.upper() for kw in ['ACTA', 'INFORME', 'MEMORIA', 'PROPUESTA'])):
                if len(texto) < 100:
                    info['title'] = texto
                    break
        text = ' '.join([p.text for p in doc.paragraphs[:10]])
        location_pattern = re.search(r'(?:CALLE|AVENIDA|AVDA|C/|AVENIDA DE)[^\n]{0,150}', text, re.IGNORECASE)
        if location_pattern:
            info['location'] = location_pattern.group(0).strip()
    return info

def apply_agoin_format_final(input_doc, project_title, project_location, is_text_only=False):
    """Aplica formato AGOIN completo"""
    output_doc = Document()

    for section in output_doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)

        # ENCABEZADO - TODO ALINEADO DERECHA
        header = section.header
        header.is_linked_to_previous = False
        for para in header.paragraphs:
            para.clear()

        # Título verde - derecha
        header_title = header.paragraphs[0]
        add_green_header_paragraph(header_title, project_title if project_title else "[TÍTULO DEL DOCUMENTO]")

        # Dirección - derecha
        header_location = header.add_paragraph()
        header_location.text = project_location if project_location else "[DIRECCIÓN DEL PROYECTO]"
        header_location.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header_location.runs:
            run.font.name = 'Century Gothic'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0, 0, 0)

        header.add_paragraph()

        # Empresa - derecha
        header_company = header.add_paragraph()
        header_company.text = "ARQUITECTURA Y GESTIÓN DE OPERACIONES INMOBILIARIAS, S.L.P."
        header_company.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header_company.runs:
            run.font.name = 'Century Gothic'
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0, 0, 0)

        # Contacto - derecha
        header_contact = header.add_paragraph()
        header_contact.text = "AVDA. DE IRLANDA 21, 4º D. 45005 TOLEDO | TLFN. 925 299 300 | www.agoin.es | info@agoin.es"
        header_contact.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header_contact.runs:
            run.font.name = 'Century Gothic'
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(51, 51, 51)

        # PIE CON LOGO IZQUIERDA - CORREGIDO
        footer = section.footer
        footer.is_linked_to_previous = False
        for para in footer.paragraphs:
            para.clear()

        # Crear tabla SIN parámetro width (causaba el error)
        footer_table = footer.paragraphs[0]._element
        tbl = OxmlElement('w:tbl')

        # Propiedades de tabla
        tblPr = OxmlElement('w:tblPr')
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), '5000')
        tblW.set(qn('w:type'), 'pct')
        tblPr.append(tblW)

        # Bordes invisibles
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        tbl.append(tblPr)

        # Grid
        tblGrid = OxmlElement('w:tblGrid')
        gridCol1 = OxmlElement('w:gridCol')
        gridCol1.set(qn('w:w'), '1500')
        gridCol2 = OxmlElement('w:gridCol')
        gridCol2.set(qn('w:w'), '7500')
        tblGrid.append(gridCol1)
        tblGrid.append(gridCol2)
        tbl.append(tblGrid)

        # Fila
        tr = OxmlElement('w:tr')

        # Celda 1: Logo
        tc1 = OxmlElement('w:tc')
        tcPr1 = OxmlElement('w:tcPr')
        tcW1 = OxmlElement('w:tcW')
        tcW1.set(qn('w:w'), '1500')
        tcW1.set(qn('w:type'), 'dxa')
        tcPr1.append(tcW1)
        tc1.append(tcPr1)

        p1 = OxmlElement('w:p')
        pPr1 = OxmlElement('w:pPr')
        jc1 = OxmlElement('w:jc')
        jc1.set(qn('w:val'), 'left')
        pPr1.append(jc1)
        p1.append(pPr1)

        # Añadir logo
        try:
            r1 = OxmlElement('w:r')
            logo_bytes = base64.b64decode(LOGO_BASE64)
            from docx.oxml.shared import OxmlElement
            from docx.oxml import parse_xml

            # Crear párrafo temporal para añadir imagen
            temp_para = footer.add_paragraph()
            temp_run = temp_para.add_run()
            logo_stream = BytesIO(logo_bytes)
            temp_run.add_picture(logo_stream, width=Inches(0.45))

            # Mover imagen a la celda
            drawing = temp_para._element.xpath('.//w:drawing')[0]
            r1.append(drawing)
            p1.append(r1)
            temp_para._element.getparent().remove(temp_para._element)
        except:
            pass

        tc1.append(p1)
        tr.append(tc1)

        # Celda 2: Texto
        tc2 = OxmlElement('w:tc')
        tcPr2 = OxmlElement('w:tcPr')
        tcW2 = OxmlElement('w:tcW')
        tcW2.set(qn('w:w'), '7500')
        tcW2.set(qn('w:type'), 'dxa')
        tcPr2.append(tcW2)
        tc2.append(tcPr2)

        # Párrafo 1
        p2 = OxmlElement('w:p')
        pPr2 = OxmlElement('w:pPr')
        jc2 = OxmlElement('w:jc')
        jc2.set(qn('w:val'), 'left')
        pPr2.append(jc2)
        p2.append(pPr2)

        r2 = OxmlElement('w:r')
        rPr2 = OxmlElement('w:rPr')
        rFonts2 = OxmlElement('w:rFonts')
        rFonts2.set(qn('w:ascii'), 'Century Gothic')
        rFonts2.set(qn('w:hAnsi'), 'Century Gothic')
        sz2 = OxmlElement('w:sz')
        sz2.set(qn('w:val'), '16')
        rPr2.append(rFonts2)
        rPr2.append(sz2)
        r2.append(rPr2)

        t2 = OxmlElement('w:t')
        t2.text = "ARQUITECTURA Y GESTIÓN DE OPERACIONES INMOBILIARIAS, S.L.P."
        r2.append(t2)
        p2.append(r2)
        tc2.append(p2)

        # Párrafo 2
        p3 = OxmlElement('w:p')
        pPr3 = OxmlElement('w:pPr')
        jc3 = OxmlElement('w:jc')
        jc3.set(qn('w:val'), 'left')
        pPr3.append(jc3)
        p3.append(pPr3)

        r3 = OxmlElement('w:r')
        rPr3 = OxmlElement('w:rPr')
        rFonts3 = OxmlElement('w:rFonts')
        rFonts3.set(qn('w:ascii'), 'Century Gothic')
        rFonts3.set(qn('w:hAnsi'), 'Century Gothic')
        sz3 = OxmlElement('w:sz')
        sz3.set(qn('w:val'), '16')
        color3 = OxmlElement('w:color')
        color3.set(qn('w:val'), '666666')
        rPr3.append(rFonts3)
        rPr3.append(sz3)
        rPr3.append(color3)
        r3.append(rPr3)

        t3 = OxmlElement('w:t')
        t3.text = "AVDA. DE IRLANDA 21, 4º D. 45005 TOLEDO | TLFN. 925 299 300 | www.agoin.es | info@agoin.es"
        r3.append(t3)
        p3.append(r3)
        tc2.append(p3)

        tr.append(tc2)
        tbl.append(tr)

        footer.paragraphs[0]._element.addprevious(tbl)

    # CONTENIDO
    if is_text_only:
        for line in input_doc.split('\n'):
            line = line.strip()
            if not line:
                continue
            new_para = output_doc.add_paragraph()
            if line.isupper() and len(line) < 100:
                add_green_header_paragraph(new_para, line)
            elif re.match(r'^\d+[.-]', line):
                run = new_para.add_run(line)
                run.font.name = 'Century Gothic'
                run.font.size = Pt(10)
                run.font.bold = True
                new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                new_para.paragraph_format.space_before = Pt(12)
                new_para.paragraph_format.space_after = Pt(6)
            else:
                run = new_para.add_run(line)
                run.font.name = 'Century Gothic'
                run.font.size = Pt(10)
                new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                new_para.paragraph_format.line_spacing = 1.15
    else:
        for para in input_doc.paragraphs:
            if not para.text.strip():
                continue
            new_para = output_doc.add_paragraph()
            texto = para.text.strip()
            es_verde = (texto.isupper() and len(texto) < 100) or para.style.name == 'Heading 1'
            es_negrita = re.match(r'^\d+[.-]', texto) or para.style.name.startswith('Heading')

            if es_verde:
                add_green_header_paragraph(new_para, texto)
            elif es_negrita:
                run = new_para.add_run(texto)
                run.font.name = 'Century Gothic'
                run.font.size = Pt(10)
                run.font.bold = True
                new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                new_para.paragraph_format.space_before = Pt(12)
                new_para.paragraph_format.space_after = Pt(6)
            else:
                for run in para.runs:
                    new_run = new_para.add_run(run.text)
                    new_run.font.name = 'Century Gothic'
                    new_run.font.size = Pt(10)
                    if run.bold:
                        new_run.bold = True
                    if run.italic:
                        new_run.italic = True
                new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                new_para.paragraph_format.line_spacing = 1.15

        for table in input_doc.tables:
            new_table = output_doc.add_table(rows=len(table.rows), cols=len(table.columns))
            new_table.style = 'Table Grid'
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    new_table.rows[i].cells[j].text = cell.text
                    for paragraph in new_table.rows[i].cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Century Gothic'
                            run.font.size = Pt(9)

    return output_doc

# INTERFAZ
col1, col2 = st.columns([2, 1])
with col1:
    st.markdown("### 📤 Subir Documento")
    uploaded_file = st.file_uploader("DOCX o TXT", type=['docx', 'txt'])
with col2:
    st.markdown('<div class="info-box"><h4>✅ Formato AGOIN</h4><p>✓ Títulos verdes derecha</p><p>✓ Logo pie izquierda</p><p>✓ Century Gothic</p></div>', unsafe_allow_html=True)

if uploaded_file:
    try:
        ext = uploaded_file.name.split('.')[-1].lower()
        doc = uploaded_file.read().decode('utf-8', errors='ignore') if ext == 'txt' else Document(uploaded_file)
        is_text = ext == 'txt'
        st.success("✅ Cargado")
        info = extract_project_info(doc)

        col_a, col_b = st.columns(2)
        with col_a:
            project_title = st.text_input("Título", value=info['title'])
        with col_b:
            project_location = st.text_area("Dirección", value=info['location'], height=80)

        if st.button("✨ Convertir al Formato AGOIN", use_container_width=True):
            with st.spinner("🔄 Aplicando formato..."):
                try:
                    output_doc = apply_agoin_format_final(doc, project_title, project_location, is_text)
                    buffer = io.BytesIO()
                    output_doc.save(buffer)
                    buffer.seek(0)
                    st.success("✅ ¡Documento formateado correctamente!")
                    st.download_button("📥 Descargar Documento AGOIN", buffer, f"AGOIN_{uploaded_file.name.rsplit('.', 1)[0]}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                except Exception as e:
                    st.error(f"❌ Error: {str(e)}")
    except Exception as e:
        st.error(f"❌ Error al cargar: {str(e)}")
else:
    st.markdown('<div class="info-box" style="text-align: center; padding: 3rem;"><h3 style="color: #1a5c4d;">👆 Sube un documento DOCX o TXT</h3></div>', unsafe_allow_html=True)

st.markdown("---")
st.markdown('<div style="text-align: center; color: #666; padding: 2rem;"><p style="font-weight: 600; color: #1a5c4d;">AGOIN Formateador Corporativo v5.0</p><p>© 2025 AGOIN S.L.P.</p></div>', unsafe_allow_html=True)
